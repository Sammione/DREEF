from fastapi import FastAPI, UploadFile, File, HTTPException
from fastapi.responses import FileResponse
from fastapi.middleware.cors import CORSMiddleware
import pandas as pd
import openai
from docx import Document
from docx.shared import Pt
import os
from dotenv import load_dotenv
import io
import json
import re
import uuid
from pydantic import BaseModel

load_dotenv()

session_mappings = {}

openai.api_key = os.getenv("OPENAI_API_KEY")

app = FastAPI()

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Robust path finding
BASE_DIR = os.path.dirname(os.path.abspath(__file__))
# We use the COPY as the template since it has the [PLACEHOLDERS]
TEMPLATE_FILENAME = "PUE report templat- Copy.docx"
TEMPLATE_PATH = os.path.join(BASE_DIR, TEMPLATE_FILENAME)

def fill_table(table, data_rows):
    """
    Fills a table after clearing all rows except the header.
    """
    # Delete all rows except the header
    while len(table.rows) > 1:
        tbl = table._tbl
        tr = table.rows[-1]._tr
        tbl.remove(tr)

    if not data_rows:
        return
    
    for row_data in data_rows:
        new_row = table.add_row()
        # Fill cells
        for j, val in enumerate(row_data):
            if j < len(new_row.cells):
                # Ensure it's a string and handle N/A
                text_val = str(val).strip() if (val is not None and str(val).strip().lower() != 'nan') else "N/A"
                new_row.cells[j].text = text_val

def docx_replace_placeholders(doc, replacements):
    """
    Replace text placeholders [KEY] in the doc.
    """
    # Create a normalized mapping to handle case sensitivity
    norm_replacements = {str(k).upper().strip(): v for k, v in replacements.items()}
    
    def process_element(element):
        if not element.text:
            return
        # Find anything in brackets [LIKE THIS]
        matches = re.findall(r'\[([^\]]+)\]', element.text)
        for match in matches:
            match_upper = match.upper().strip()
            if match_upper in norm_replacements:
                val = str(norm_replacements[match_upper])
                element.text = element.text.replace(f"[{match}]", val)
            else:
                # If it's a placeholder but we have no data, clear it or set to N/A
                element.text = element.text.replace(f"[{match}]", "N/A")

    for p in doc.paragraphs:
        process_element(p)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    process_element(p)

def apply_mapping_to_template(mapping):
    doc = Document(TEMPLATE_PATH)
    
    placeholders = mapping.get("placeholders", {})
    new_name = str(placeholders.get("COMMUNITY NAME") or placeholders.get("Community Name") or "The Community")
    new_desc = str(placeholders.get("COMMUNITY_DESCRIPTION") or f"{new_name} is a newly evaluated community for PUE projects.")
    
    STALE_FRAGMENTS = [
        "Agadagba", "Ikale people", "Oba Michael", "Larogbo", "Ijaw", "Ado", "Hausa", 
        "20 years", "nearly 20 years", "coastal community", "sea, linking it", "marine", 
        "Larogbo of Lagos", "Akah", "Ondo"
    ]
    
    for p in doc.paragraphs:
        text = p.text.strip()
        if any(f.lower() in text.lower() for f in STALE_FRAGMENTS):
            if len(text) > 80:
                p.text = new_desc
            else:
                for f in STALE_FRAGMENTS:
                    if f.lower() in text.lower():
                        p.text = text.replace(f, new_name)
        
        docx_replace_placeholders(doc, placeholders)

    tables_data = mapping.get("tables", {})
    for i, table in enumerate(doc.tables):
        i_str = str(i)
        if i_str in tables_data:
            fill_table(table, tables_data[i_str])
        else:
            if len(table.rows) > 1:
                fill_table(table, [])

    clean_name = re.sub(r'[\\/*?:"<>|]', "", new_name).strip()
    output_name = f"Report_{clean_name}.docx"
    # Use /tmp for Vercel serverless environment compatibility (read-only filesystem)
    output_path = os.path.join("/tmp", output_name)
    doc.save(output_path)
    return output_path, output_name

class RevisionRequest(BaseModel):
    session_id: str
    instruction: str

@app.post("/api/revise_report")
async def revise_report(req: RevisionRequest):
    try:
        if req.session_id not in session_mappings:
            raise HTTPException(status_code=404, detail="Session not found")
            
        current_mapping = session_mappings[req.session_id]
        
        prompt = f"""
You are a senior report engine. 
CURRENT JSON MAPPING OF THE REPORT:
{json.dumps(current_mapping)}

USER INSTRUCTION:
"{req.instruction}"

Return the FULL, UPDATED JSON mapping incorporating the user's modifications.
You must return the EXACT same structure, just editing the requested values.
Keep the rest unchanged!
OUTPUT FORMAT MUST BE VALID JSON ONLY.
"""
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You return JSON mapping matching the provided structure."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=4096
        )
        
        new_mapping = json.loads(response.choices[0].message.content)
        session_mappings[req.session_id] = new_mapping
        
        output_path, output_name = apply_mapping_to_template(new_mapping)
        return FileResponse(output_path, media_type='application/vnd.wordprocessingml.document', filename=output_name)
    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

@app.post("/api/generate_report")
async def generate_report(file: UploadFile = File(...)):
    try:
        print(f"Received file: {file.filename}")
        contents = await file.read()
        
        # Read Excel
        try:
            # Smart find header - check first 30 rows
            best_header_idx = 0
            max_keywords = -1
            
            temp_df_raw = pd.read_excel(io.BytesIO(contents), header=None, nrows=30)
            for idx, row in temp_df_raw.iterrows():
                row_str = " ".join([str(x).lower() for x in row if pd.notna(x)])
                keywords = ['community', 'state', 'lga', 'name', 'coordinate', 'latitude', 'people interviewed', 'gender', 'crop', 'populati', 'household']
                match_count = sum(1 for k in keywords if k in row_str)
                if match_count > max_keywords:
                    max_keywords = match_count
                    best_header_idx = idx
            
            print(f"Detected header at row {best_header_idx} with score {max_keywords}")
            df = pd.read_excel(io.BytesIO(contents), header=best_header_idx)
            df = df.dropna(how='all', axis=1).fillna("N/A")
            
            # GENERATE COMPREHENSIVE DATA SUMMARY FOR AI
            summary = []
            summary.append(f"### FULL EXCEL DATA SUMMARY (Total Rows: {len(df)})")
            summary.append(f"Columns Found: {', '.join(df.columns)}")
            
            # Add categorical summary (Counts)
            for col in df.columns:
                col_lower = str(col).lower()
                # Skip numeric/date columns for counts
                if any(x in col_lower for x in ['gender', 'ethnic', 'crop', 'activity', 'source', 'machine', 'income', 'people']):
                    counts = df[col].astype(str).value_counts().head(15).to_dict()
                    summary.append(f"Counts for '{col}': {json.dumps(counts)}")
                elif any(x in col_lower for x in ['population', 'household', 'total', 'count', 'number']):
                    # Summable numeric
                    try:
                        num_sum = pd.to_numeric(df[col], errors='coerce').sum()
                        summary.append(f"Total sum for '{col}': {num_sum}")
                    except:
                        pass
            
            # Send sample rows (representative)
            data_context = df.head(50).to_markdown(index=False)
            if len(df) > 70:
                data_context += "\n\n... (Rows Skipped) ...\n\n" + df.tail(20).to_markdown(index=False)
            
            final_data_summary = "\n".join(summary) + "\n\n### REPRESENTATIVE DATA SAMPLE:\n" + data_context
            
        except Exception as e:
            import traceback
            traceback.print_exc()
            raise HTTPException(status_code=400, detail=f"Excel Parsing Error: {str(e)}")

        # Extract Template Info
        doc = Document(TEMPLATE_PATH)
        table_context = []
        for i, table in enumerate(doc.tables):
            headers = [cell.text.strip() for cell in table.rows[0].cells] if table.rows else []
            table_context.append({"index": i, "headers": headers})

        # Prompt AI for JSON mapping
        prompt = f"""
You are a senior report engine. You must convert field survey data into a clean report.

### THE DATA:
\"\"\"
{final_data_summary}
\"\"\"

### REPORT TABLES:
{json.dumps(table_context[:65], indent=2)}

### RIGID INSTRUCTIONS:
1. **Remove Old Context**: The template contains stale descriptions of "Agadagba" (Ikale people, Oba Michael, coastal, lagoon, no power for 20 years). You MUST IGNORE these.
2. **Current Overview**: Provide a NEW community overview in the "COMMUNITY_DESCRIPTION" placeholder. Use ONLY facts from Excel (like Location name, State, LGA, Crop types, and survey counts). Make it professional.
3. **Data Mapping**: Use the survey counts I provided in the SUMMARY to fill out demographic and SME tables (gender balance, community sizes, counts of agro-processors).
4. **No Hallucinations**: If you don't have details about a specific project capacity, use the values from the provided Excel summary or set to "N/A".

### OUTPUT FORMAT:
{{
  "placeholders": {{
    "COMMUNITY NAME": "Lagos",
    "COMMUNITY_DESCRIPTION": "Lagos is a community in Mushin, Lagos State, with a population focusing on agriculture and SME trading...",
    "Date": "March 2024",
    "Name of state": "Lagos",
    "Name of LGA": "Mushin"
  }},
  "tables": {{
     "0": [["Lagos", "Mushin", "10.8, 7.5", "Lagos"]],
     "1": [["Households", "350"], ["Processors", "18"], ...]
  }}
}}
"""

        print("Requesting AI for mapping...")
        response = openai.chat.completions.create(
            model="gpt-4o-mini",
            response_format={"type": "json_object"},
            messages=[
                {"role": "system", "content": "You are a professional report analyst. You extract and return survey data in JSON format only, avoiding all template residue like 'Agadagba' or 'Oba Michael'."},
                {"role": "user", "content": prompt}
            ],
            temperature=0,
            max_tokens=4096
        )
        
        mapping = json.loads(response.choices[0].message.content)
        print("Mapping received.")

        session_id = str(uuid.uuid4())
        session_mappings[session_id] = mapping
        
        output_path, output_name = apply_mapping_to_template(mapping)
        
        return FileResponse(output_path, media_type='application/vnd.wordprocessingml.document', filename=output_name, headers={"X-Session-ID": session_id, "Access-Control-Expose-Headers": "X-Session-ID"})

    except Exception as e:
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(app, host="0.0.0.0", port=8000)
